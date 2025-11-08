#!/usr/bin/env python3
"""
Convert markdown guides to .docx format
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                # Add light gray shading
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # Check if next line is not a table line
            if i >= len(lines) or '|' not in lines[i]:
                # Create table
                create_table_from_markdown(doc, table_lines)
                in_table = False
                table_lines = []
            continue

        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=0)
                p.runs[0].font.size = Pt(24)
                p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            elif level == 2:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            elif level == 3:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            else:
                p = doc.add_heading(text, level=3)
                p.runs[0].font.size = Pt(12)

            i += 1
            continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Handle checkboxes
        if line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
            checked = '[x]' in line
            text = line.strip().replace('- [ ]', '').replace('- [x]', '').strip()
            symbol = '☑' if checked else '☐'
            p = doc.add_paragraph(f'{symbol} {text}')
            i += 1
            continue

        # Regular paragraph
        text = process_inline_formatting(line.strip())
        if text:
            p = doc.add_paragraph()
            add_formatted_text(p, text)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Created: {docx_file}")

def process_inline_formatting(text):
    """Process bold, italic, code, and other inline formatting"""
    # Don't process if it's just formatting markers
    return text

def add_formatted_text(paragraph, text):
    """Add text with inline formatting to paragraph"""
    # Handle inline code
    parts = re.split(r'(`[^`]+`)', text)

    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Inline code
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            # Handle bold and italic
            process_bold_italic(paragraph, part)

def process_bold_italic(paragraph, text):
    """Process bold and italic markdown"""
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Check for italic
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    # Italic text
                    italic_text = ipart[1:-1]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                else:
                    # Regular text
                    if ipart:
                        paragraph.add_run(ipart)

def create_table_from_markdown(doc, table_lines):
    """Create a Word table from markdown table lines"""
    # Remove empty lines and separator lines
    cleaned_lines = []
    for line in table_lines:
        line = line.strip()
        if line and not re.match(r'^\|[\s\-:]+\|$', line):
            cleaned_lines.append(line)

    if not cleaned_lines:
        return

    # Parse rows
    rows = []
    for line in cleaned_lines:
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last cells from leading/trailing |
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Create table
    num_cols = len(rows[0])
    num_rows = len(rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()  # Add space after table

if __name__ == '__main__':
    # Convert all guides
    parse_markdown_to_docx('HAND_FIXING_GUIDE.md', 'HAND_FIXING_GUIDE.docx')
    parse_markdown_to_docx('FIXING_HAND_ANATOMY.md', 'FIXING_HAND_ANATOMY.docx')
    parse_markdown_to_docx('INSTALLATION_TROUBLESHOOTING.md', 'INSTALLATION_TROUBLESHOOTING.docx')
    print("\nConversion complete!")
